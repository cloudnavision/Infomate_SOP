"""
Build the SOP DOCX template from scratch using python-docx.
Matches the Aged Debtor Process document structure:
  - Cover page with title + metadata table
  - Table of Contents (numbered sections, dot leaders)
  - Pre-sections (Heading 2, numbered)
  - Process Map section (Heading 1, numbered)
  - Detailed Procedure (Heading 1, numbered) with steps (Heading 3)
  - Post-sections (Heading 2, numbered)

Jinja2 / docxtpl tags are embedded as plain text in the document.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor, Cm


# ── Brand colours (matching Aged Debtor / Starboard Hotels palette) ──────────
ORANGE   = RGBColor(0xE8, 0x5C, 0x1A)   # #E85C1A
DARK     = RGBColor(0x1A, 0x1A, 0x2E)   # near-black heading
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = RGBColor(0xF8, 0xF9, 0xFA)   # table alt row
BORDER   = RGBColor(0xD1, 0xD5, 0xDB)   # table border

TEMPLATE_PATH = Path("/data/templates/sop_template.docx")
_VERSION_PATH = TEMPLATE_PATH.with_suffix(".version")
_TEMPLATE_VERSION = "4"  # increment when template structure changes


def _set_run_font(run, size_pt: float, bold=False, italic=False, color=None):
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color


def _set_para_spacing(para, before_pt=0, after_pt=6, line_rule=None, line_val=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(int(before_pt * 20)))
    spacing.set(qn("w:after"), str(int(after_pt * 20)))
    if line_rule and line_val:
        spacing.set(qn("w:lineRule"), line_rule)
        spacing.set(qn("w:line"), str(int(line_val * 240)))
    pPr.append(spacing)


def _para_shade(para, fill_hex: str):
    """Set paragraph background shading."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def _set_cell_bg(cell, fill_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _table_borders(tbl):
    """Add thin borders to all table cells."""
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "D1D5DB")
        borders.append(b)
    tblPr.append(borders)


def _add_toc_entry(doc, num: str, title_tag: str, level: int = 0):
    """
    Add a single TOC line with a dot leader tab.
    level 0 = main section  (bold, numbered, no indent)
    level 1 = sub-item      (indented, no number)
    Uses static indentation values — no Jinja2 in XML attributes.
    """
    p = doc.add_paragraph()
    p.style = "Normal"

    pPr = p._p.get_or_add_pPr()

    # Static indent based on level (no Jinja2 expressions in XML attributes)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360" if level > 0 else "0")
    ind.set(qn("w:hanging"), "0")
    pPr.append(ind)

    # Tab stop: right-aligned dot-leader at 14 cm
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), "7920")   # 14 cm ≈ 7920 twips
    tabs.append(tab)
    pPr.append(tabs)
    _set_para_spacing(p, before_pt=2, after_pt=2)

    if num:
        r_num = p.add_run(f"{num}  ")
        _set_run_font(r_num, 11, bold=True, color=DARK)

    r_title = p.add_run(title_tag)
    _set_run_font(r_title, 10 if level == 0 else 9.5, bold=(level == 0), color=DARK)

    # Dot leader tab + page ref placeholder
    r_tab = p.add_run("\t")
    _set_run_font(r_tab, 10)


def _ctrl_para(doc, tag: str):
    """Add a Jinja2 control tag paragraph (for, endfor, if, else, endif)."""
    p = doc.add_paragraph(tag)
    p.style = "Normal"
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)


def build(force: bool = False):
    if not force and TEMPLATE_PATH.exists():
        # Skip rebuild only if version matches
        current = _VERSION_PATH.read_text().strip() if _VERSION_PATH.exists() else ""
        if current == _TEMPLATE_VERSION:
            return

    TEMPLATE_PATH.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # ── Page margins (matching A4 professional) ───────────────────────────────
    for sec in doc.sections:
        sec.page_width  = Cm(21)
        sec.page_height = Cm(29.7)
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(2.0)

    # ── Default paragraph style ───────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # ── COVER PAGE ────────────────────────────────────────────────────────────
    # Orange header bar
    p_cover_bar = doc.add_paragraph()
    _para_shade(p_cover_bar, "E85C1A")
    _set_para_spacing(p_cover_bar, before_pt=0, after_pt=0)
    p_cover_bar.paragraph_format.space_before = Pt(0)
    r = p_cover_bar.add_run("  STANDARD OPERATING PROCEDURE")
    _set_run_font(r, 13, bold=True, color=WHITE)

    # Title
    p_title = doc.add_paragraph()
    _set_para_spacing(p_title, before_pt=24, after_pt=6)
    r = p_title.add_run("{{ sop_title }}")
    _set_run_font(r, 22, bold=True, color=DARK)

    # Sub-line: process name
    p_proc = doc.add_paragraph()
    _set_para_spacing(p_proc, before_pt=0, after_pt=18)
    r = p_proc.add_run("{{ process_name }}")
    _set_run_font(r, 13, italic=True, color=ORANGE)

    # Metadata table
    tbl_meta = doc.add_table(rows=5, cols=2)
    tbl_meta.style = "Table Grid"
    _table_borders(tbl_meta)
    tbl_meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta_rows = [
        ("Client",        "{{ client_name }}"),
        ("Process",       "{{ process_name }}"),
        ("Meeting Date",  "{{ meeting_date }}"),
        ("Generated",     "{{ generated_date }}"),
        ("Total Steps",   "{{ step_count }}"),
    ]
    for i, (label, value) in enumerate(meta_rows):
        row = tbl_meta.rows[i]
        row.cells[0].width = Inches(1.8)
        row.cells[1].width = Inches(3.5)
        _set_cell_bg(row.cells[0], "F3F4F6")
        p0 = row.cells[0].paragraphs[0]
        r0 = p0.add_run(label)
        _set_run_font(r0, 10, bold=True, color=DARK)
        p1 = row.cells[1].paragraphs[0]
        r1 = p1.add_run(value)
        _set_run_font(r1, 10, color=DARK)

    # Page break before TOC
    doc.add_page_break()

    # ── TABLE OF CONTENTS ─────────────────────────────────────────────────────
    p_toc_head = doc.add_heading("Table of Contents", level=1)
    _set_para_spacing(p_toc_head, before_pt=0, after_pt=12)
    for run in p_toc_head.runs:
        _set_run_font(run, 16, bold=True, color=ORANGE)

    # Orange underline bar
    p_bar = doc.add_paragraph()
    _para_shade(p_bar, "E85C1A")
    _set_para_spacing(p_bar, before_pt=0, after_pt=8)
    r = p_bar.add_run(" ")
    _set_run_font(r, 4)

    # Jinja2 loop — use if/else to select main vs sub-item paragraph style.
    # This avoids putting Jinja2 expressions inside XML attributes (unreliable).
    _ctrl_para(doc, "{%- for entry in toc_entries %}")
    _ctrl_para(doc, "{%- if not entry.is_sub %}")
    _add_toc_entry(doc, num="{{ entry.num }}", title_tag="{{ entry.title }}", level=0)
    _ctrl_para(doc, "{%- else %}")
    _add_toc_entry(doc, num="{{ entry.num }}", title_tag="{{ entry.title }}", level=1)
    _ctrl_para(doc, "{%- endif %}")
    _ctrl_para(doc, "{%- endfor %}")

    doc.add_page_break()

    # ── PRE-SECTIONS ─────────────────────────────────────────────────────────
    _ctrl_para(doc, "{%- for section in sections_pre %}")

    h2 = doc.add_heading("{{ section.num }}  {{ section.section_title }}", level=2)
    for run in h2.runs:
        _set_run_font(run, 13, bold=True, color=ORANGE)
    _set_para_spacing(h2, before_pt=18, after_pt=6)

    p_content = doc.add_paragraph("{{r section.content_text }}")
    p_content.style = "Normal"

    _ctrl_para(doc, "{%- endfor %}")

    doc.add_page_break()

    # ── PROCESS MAP ───────────────────────────────────────────────────────────
    h1_pm = doc.add_heading("{{ pm_section_num }}  Process Map", level=1)
    for run in h1_pm.runs:
        _set_run_font(run, 15, bold=True, color=ORANGE)
    _set_para_spacing(h1_pm, before_pt=12, after_pt=8)

    p_pm = doc.add_paragraph("{%- if process_map %}{{ process_map }}{%- else %}(Process map not configured — use the Process Map tab to build one){%- endif %}")
    p_pm.style = "Normal"

    doc.add_page_break()

    # ── DETAILED PROCEDURE ────────────────────────────────────────────────────
    h1_dp = doc.add_heading("{{ dp_section_num }}  Detailed Procedure", level=1)
    for run in h1_dp.runs:
        _set_run_font(run, 15, bold=True, color=ORANGE)
    _set_para_spacing(h1_dp, before_pt=12, after_pt=8)

    _ctrl_para(doc, "{%- for step in steps %}")

    # Step heading
    h3 = doc.add_heading("Step {{ step.sequence }}: {{ step.title }}", level=3)
    for run in h3.runs:
        _set_run_font(run, 12, bold=True, color=DARK)
    _set_para_spacing(h3, before_pt=16, after_pt=4)

    # Description
    p_desc = doc.add_paragraph("{{ step.description | default('') }}")
    p_desc.style = "Normal"
    _set_para_spacing(p_desc, before_pt=0, after_pt=4)

    # Sub-steps
    _ctrl_para(doc, "{%- for sub in step.sub_steps %}")
    p_sub = doc.add_paragraph("{{ sub }}")
    p_sub.style = "List Bullet"
    _set_para_spacing(p_sub, before_pt=0, after_pt=2)
    _ctrl_para(doc, "{%- endfor %}")

    # Screenshot
    p_ss = doc.add_paragraph("{%- if step.screenshot %}{{ step.screenshot }}{%- endif %}")
    p_ss.style = "Normal"
    _set_para_spacing(p_ss, before_pt=4, after_pt=4)

    # Callouts
    _ctrl_para(doc, "{%- if step.callouts %}")

    p_callout_head = doc.add_paragraph("Callout References")
    p_callout_head.style = "Normal"
    for run in p_callout_head.runs:
        _set_run_font(run, 10, bold=True, italic=True, color=DARK)
    _set_para_spacing(p_callout_head, before_pt=4, after_pt=2)

    _ctrl_para(doc, "{%- for callout in step.callouts %}")
    p_cl = doc.add_paragraph("{{ callout.callout_number }}. {{ callout.label }}")
    p_cl.style = "List Number"
    _set_para_spacing(p_cl, before_pt=0, after_pt=2)
    _ctrl_para(doc, "{%- endfor %}")
    _ctrl_para(doc, "{%- endif %}")

    # Step separator
    p_sep = doc.add_paragraph()
    p_sep.style = "Normal"
    _para_shade(p_sep, "F3F4F6")
    _set_para_spacing(p_sep, before_pt=12, after_pt=0)
    r = p_sep.add_run(" ")
    _set_run_font(r, 3)

    _ctrl_para(doc, "{%- endfor %}")

    doc.add_page_break()

    # ── POST-SECTIONS ─────────────────────────────────────────────────────────
    _ctrl_para(doc, "{%- for section in sections_post %}")

    h2_post = doc.add_heading("{{ section.num }}  {{ section.section_title }}", level=2)
    for run in h2_post.runs:
        _set_run_font(run, 13, bold=True, color=ORANGE)
    _set_para_spacing(h2_post, before_pt=18, after_pt=6)

    p_post = doc.add_paragraph("{{r section.content_text }}")
    p_post.style = "Normal"

    _ctrl_para(doc, "{%- endfor %}")

    doc.save(str(TEMPLATE_PATH))
    _VERSION_PATH.write_text(_TEMPLATE_VERSION)
    print(f"Template v{_TEMPLATE_VERSION} written to {TEMPLATE_PATH}")


if __name__ == "__main__":
    build(force=True)
