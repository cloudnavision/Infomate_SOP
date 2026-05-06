"""
SOP Document Renderer
Phase 7a: docxtpl template injection + LibreOffice PDF conversion + Azure Blob upload
"""
import logging
import re
import subprocess
import tempfile
import time
from datetime import date
from pathlib import Path
from typing import Optional

import requests
from docxtpl import DocxTemplate, InlineImage, RichText
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

TEMPLATE_PATH = Path("/data/templates/sop_template.docx")
EXPORTS_DIR = Path("/data/exports")

# Regex for characters that are illegal in XML 1.0 (excludes \t \n \r which are fine)
_XML_ILLEGAL_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")


def _sanitize_text(text: str) -> str:
    """Strip XML-illegal characters that would corrupt the DOCX output."""
    if not text:
        return text
    return _XML_ILLEGAL_RE.sub("", text)


def render_sop(
    sop_id: str,
    fmt: str,                   # 'docx' or 'pdf'
    sop_data: dict,
    azure_blob_base_url: str,   # e.g. https://cnavinfsop.blob.core.windows.net/infsop
    azure_sas_token: str,
) -> dict:
    """
    Render a SOP document from the Word template.

    Returns:
        {"docx_url": str, "pdf_url": str | None}
        URLs are base Azure Blob URLs without SAS (safe for DB storage).
    """
    # Rebuild template if missing or outdated
    try:
        from app.build_template import build as build_template
        build_template(force=False)
    except Exception as exc:
        logger.warning("Template builder failed: %s", exc)

    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Template not found: {TEMPLATE_PATH}")

    export_dir = EXPORTS_DIR / sop_id
    export_dir.mkdir(parents=True, exist_ok=True)

    with tempfile.TemporaryDirectory(prefix=f"sop_render_{sop_id}_") as tmp_str:
        tmp_dir = Path(tmp_str)

        tpl = DocxTemplate(str(TEMPLATE_PATH))
        table_registry: dict[str, list] = {}
        context = _build_context(tpl, sop_data, tmp_dir, table_registry, azure_sas_token=azure_sas_token)
        tpl.render(context)

        # Save rendered docx
        docx_filename = f"sop_{sop_id}.docx"
        docx_path = export_dir / docx_filename
        tpl.save(str(docx_path))

        # Post-process: replace table placeholders with real Word tables
        if table_registry:
            _inject_tables(docx_path, table_registry)

        # Post-process: add TOC hyperlinks (bookmarks on headings + hyperlink fields in TOC)
        _inject_toc_links(docx_path)

        logger.info("Rendered DOCX: %s (%.1f KB)", docx_path, docx_path.stat().st_size / 1024)

        # Upload DOCX
        docx_blob_path = f"exports/{sop_id}/{docx_filename}"
        docx_base_url = f"{azure_blob_base_url}/{docx_blob_path}"
        _upload_blob(
            docx_path,
            f"{docx_base_url}?{azure_sas_token}",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        logger.info("Uploaded DOCX → %s", docx_blob_path)

        pdf_base_url: Optional[str] = None

        if fmt == "pdf":
            pdf_path = _convert_to_pdf(docx_path, export_dir)
            logger.info("PDF created: %s (%.1f KB)", pdf_path, pdf_path.stat().st_size / 1024)

            pdf_filename = pdf_path.name
            pdf_blob_path = f"exports/{sop_id}/{pdf_filename}"
            pdf_base_url = f"{azure_blob_base_url}/{pdf_blob_path}"
            _upload_blob(pdf_path, f"{pdf_base_url}?{azure_sas_token}", "application/pdf")
            logger.info("Uploaded PDF → %s", pdf_blob_path)

    return {"docx_url": docx_base_url, "pdf_url": pdf_base_url}


def _section_content(tpl: DocxTemplate, section: dict, table_registry: dict):
    """Return a RichText for a section. Table sections use a placeholder string."""
    content_type = str(section.get("content_type") or "text")
    if "." in content_type:
        content_type = content_type.split(".")[-1]

    text = _sanitize_text(section.get("content_text") or "")
    json_data = section.get("content_json")

    if content_type == "table" and isinstance(json_data, list) and json_data:
        rows_data = [r for r in json_data if isinstance(r, dict)]
        if rows_data:
            placeholder = f"__TBLPH_{len(table_registry):03d}__"
            table_registry[placeholder] = rows_data
            rt = RichText()
            rt.add(placeholder)
            return rt

    if content_type == "list":
        rt = RichText()
        items = json_data if isinstance(json_data, list) else ([json_data] if json_data else [text])
        for i, item in enumerate(items):
            item_str = _sanitize_text(str(item)) if item is not None else ""
            if i > 0:
                rt.xml += "<w:r><w:br/></w:r>"
            rt.add(f"•  {item_str}", font="Calibri", size=22)
        return rt

    rt = RichText()
    rt.add(text, font="Calibri", size=22)
    return rt


def _set_cell_shd(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _add_tbl_borders(tbl) -> None:
    tblPr = tbl._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "D1D5DB")
        borders.append(b)
    tblPr.append(borders)


def _inject_tables(docx_path: Path, table_registry: dict[str, list]) -> None:
    """Open the saved docx, find placeholder paragraphs, replace with Word tables."""
    from docx import Document as DocxDoc

    doc = DocxDoc(str(docx_path))
    body = doc.element.body

    for para in doc.paragraphs:
        text = para.text.strip()
        if text not in table_registry:
            continue
        rows_data = table_registry[text]
        headers = list(rows_data[0].keys())
        n_cols = len(headers)
        col_w = Inches(5.3 / n_cols)

        # Build the table
        tbl = doc.add_table(rows=1 + len(rows_data), cols=n_cols)
        _add_tbl_borders(tbl)

        # Header row — orange bg, white bold
        for i, h in enumerate(headers):
            cell = tbl.rows[0].cells[i]
            cell.width = col_w
            _set_cell_shd(cell, "E85C1A")
            run = cell.paragraphs[0].add_run(h.replace("_", " ").upper())
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9)
            run.font.name = "Calibri"

        # Data rows — alternating bg
        for r_idx, row_data in enumerate(rows_data):
            bg = "F8F9FA" if r_idx % 2 == 0 else "FFFFFF"
            for c_idx, h in enumerate(headers):
                cell = tbl.rows[r_idx + 1].cells[c_idx]
                cell.width = col_w
                _set_cell_shd(cell, bg)
                run = cell.paragraphs[0].add_run(str(row_data.get(h) or ""))
                run.font.size = Pt(9)
                run.font.name = "Calibri"

        # Move the new table element to replace the placeholder paragraph
        tbl_element = tbl._tbl
        body.remove(tbl_element)          # add_table appended to body; move it
        para._element.addprevious(tbl_element)
        para._element.getparent().remove(para._element)

    doc.save(str(docx_path))


# ── TOC Hyperlinks ────────────────────────────────────────────────────────────

def _add_para_bookmark(para, name: str, bm_id: int) -> None:
    """Insert a named bookmark at the start of a heading paragraph."""
    p = para._p
    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), str(bm_id))
    bm_start.set(qn("w:name"), name)
    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), str(bm_id))

    # Insert immediately after pPr (or at position 0 if no pPr)
    children = list(p)
    pPr = p.find(qn("w:pPr"))
    insert_pos = children.index(pPr) + 1 if pPr is not None else 0
    p.insert(insert_pos, bm_start)
    p.insert(insert_pos + 1, bm_end)


def _is_toc_para(para) -> bool:
    """Return True if paragraph has the TOC right-aligned dot-leader tab at 7920 twips."""
    pPr = para._p.find(qn("w:pPr"))
    if pPr is None:
        return False
    tabs = pPr.find(qn("w:tabs"))
    if tabs is None:
        return False
    for tab in tabs.findall(qn("w:tab")):
        if (
            tab.get(qn("w:val")) == "right"
            and tab.get(qn("w:leader")) == "dot"
            and tab.get(qn("w:pos")) == "7920"
        ):
            return True
    return False


def _add_toc_hyperlink(para, anchor: str) -> None:
    """Wrap all runs in a TOC paragraph inside a w:hyperlink element pointing to anchor."""
    p = para._p
    runs = p.findall(qn("w:r"))
    if not runs:
        return

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)

    # Insert hyperlink at position of first run, then move all runs into it
    children = list(p)
    first_pos = children.index(runs[0])
    p.insert(first_pos, hyperlink)
    for r in runs:
        p.remove(r)
        hyperlink.append(r)


def _inject_toc_links(docx_path: Path) -> None:
    """
    Post-process rendered DOCX:
    1. Add named bookmarks to all Heading 1/2/3 paragraphs.
    2. Wrap TOC entry paragraphs (dot-leader tab at 7920) in w:hyperlink elements.
    Headings and TOC entries are matched by normalising whitespace in their text.
    """
    try:
        from docx import Document as DocxDoc

        doc = DocxDoc(str(docx_path))
        all_paras = list(doc.paragraphs)
        heading_styles = {"Heading 1", "Heading 2", "Heading 3"}

        # ── Pass 1: assign bookmarks to heading paragraphs ────────────────────
        bm_id = 200  # start high to avoid conflicts with Word's auto-bookmarks
        bm_map: dict[str, str] = {}  # normalised_text → bookmark_name

        for para in all_paras:
            style = para.style.name if para.style else ""
            if style not in heading_styles:
                continue
            text = " ".join(para.text.split())
            if not text:
                continue
            bm_name = f"_soptoc{bm_id}"
            bm_map[text] = bm_name
            _add_para_bookmark(para, bm_name, bm_id)
            bm_id += 1

        if not bm_map:
            return  # no headings found — skip (e.g. empty SOP)

        # ── Pass 2: add hyperlinks to TOC paragraphs ─────────────────────────
        for para in all_paras:
            if not _is_toc_para(para):
                continue
            raw_text = para.text
            # Strip tab character and anything after it (dot leaders + empty page ref)
            title_part = raw_text.split("\t")[0]
            normalised = " ".join(title_part.split())
            if not normalised:
                continue
            bm_name = bm_map.get(normalised)
            if bm_name:
                _add_toc_hyperlink(para, bm_name)

        doc.save(str(docx_path))
        logger.info("Injected TOC hyperlinks (%d bookmarks)", len(bm_map))

    except Exception as exc:
        logger.warning("TOC link injection failed (non-fatal): %s", exc)


# ── Context builder ───────────────────────────────────────────────────────────

def _build_context(tpl: DocxTemplate, sop_data: dict, tmp_dir: Path, table_registry: dict | None = None, azure_sas_token: str = "") -> dict:
    """Build the Jinja2 context dict for docxtpl."""
    steps_raw = sop_data.get("steps", [])
    steps_ctx = []

    for step in steps_raw:
        screenshot = None
        ann_url = step.get("annotated_screenshot_url") or step.get("screenshot_url")
        if ann_url:
            screenshot = _download_inline_image(tpl, ann_url, tmp_dir, step.get("id", "unknown"))

        steps_ctx.append({
            "sequence": step.get("sequence", ""),
            "title": _sanitize_text(step.get("title") or ""),
            "description": _sanitize_text(step.get("description") or ""),
            "sub_steps": [_sanitize_text(str(s)) for s in (step.get("sub_steps") or []) if s is not None],
            "screenshot": screenshot,
            "callouts": [
                {
                    "callout_number": c.get("callout_number"),
                    "label": _sanitize_text(c.get("label") or ""),
                }
                for c in (step.get("callouts") or [])
            ],
        })

    # Split sections: display_order < 50 before procedure, >= 50 after
    all_sections = sop_data.get("sections") or []
    raw_pre  = [s for s in all_sections if (s.get("display_order") or 0) < 50]
    raw_post = [s for s in all_sections if (s.get("display_order") or 0) >= 50]

    # Assign section numbers sequentially (matching Aged Debtor TOC style)
    sec_num = 1
    sections_pre = []
    for s in raw_pre:
        sections_pre.append({
            "num": str(sec_num),
            "section_title": _sanitize_text(s.get("section_title") or ""),
            "content_text": _section_content(tpl, s, table_registry if table_registry is not None else {}),
        })
        sec_num += 1

    pm_section_num = str(sec_num); sec_num += 1
    dp_section_num = str(sec_num); sec_num += 1

    sections_post = []
    for s in raw_post:
        sections_post.append({
            "num": str(sec_num),
            "section_title": _sanitize_text(s.get("section_title") or ""),
            "content_text": _section_content(tpl, s, table_registry if table_registry is not None else {}),
        })
        sec_num += 1

    pm_config = sop_data.get("process_map_config")
    confirmed_url = pm_config.get("confirmed_url") if pm_config else None

    if confirmed_url:
        process_map = _download_confirmed_map(tpl, confirmed_url, tmp_dir, sas_token=azure_sas_token)
        if process_map is None:
            process_map = (
                _generate_swimlane_map(tpl, pm_config, steps_raw, tmp_dir)
                if pm_config and pm_config.get("lanes") and pm_config.get("assignments")
                else _generate_process_map(tpl, steps_raw, tmp_dir)
            )
    elif pm_config and pm_config.get("lanes") and pm_config.get("assignments"):
        process_map = _generate_swimlane_map(tpl, pm_config, steps_raw, tmp_dir)
    else:
        process_map = _generate_process_map(tpl, steps_raw, tmp_dir)
    today = date.today().strftime("%d %b %Y")

    # ── Build TOC entries ─────────────────────────────────────────────────────
    # is_sub=False → main entry (Heading 1/2, with section number, no indent)
    # is_sub=True  → sub-item (Heading 3 step, indented, no section number)
    toc_entries = []
    for s in sections_pre:
        toc_entries.append({"num": s["num"], "title": s["section_title"], "is_sub": False})

    toc_entries.append({"num": pm_section_num, "title": "Process Map", "is_sub": False})
    toc_entries.append({"num": dp_section_num, "title": "Detailed Procedure", "is_sub": False})

    for s in sections_post:
        toc_entries.append({"num": s["num"], "title": s["section_title"], "is_sub": False})

    return {
        "sop_title": _sanitize_text(sop_data.get("sop_title") or ""),
        "client_name": _sanitize_text(sop_data.get("client_name") or ""),
        "process_name": _sanitize_text(sop_data.get("process_name") or ""),
        "meeting_date": _sanitize_text(sop_data.get("meeting_date") or ""),
        "generated_date": today,
        "step_count": sop_data.get("step_count", len(steps_raw)),
        "steps": steps_ctx,
        "sections_pre": sections_pre,
        "sections_post": sections_post,
        "pm_section_num": pm_section_num,
        "dp_section_num": dp_section_num,
        "process_map": process_map,
        "toc_entries": toc_entries,
    }


def _download_inline_image(
    tpl: DocxTemplate,
    url: str,
    tmp_dir: Path,
    step_id: str,
) -> Optional[InlineImage]:
    """
    Download a screenshot, resize to max 1400 px wide, and save as JPEG.
    Keeps DOCX size small and render time fast.
    """
    try:
        from PIL import Image as PILImage
        import io

        resp = requests.get(url, timeout=30)
        resp.raise_for_status()

        img = PILImage.open(io.BytesIO(resp.content)).convert("RGB")
        max_w = 1400
        if img.width > max_w:
            ratio = max_w / img.width
            img = img.resize((max_w, int(img.height * ratio)), PILImage.LANCZOS)

        img_path = tmp_dir / f"screenshot_{step_id}.jpg"
        img.save(str(img_path), "JPEG", quality=85, optimize=True)

        return InlineImage(tpl, str(img_path), width=Inches(5.5))
    except Exception as exc:
        logger.warning("Could not download screenshot for step %s: %s", step_id, exc)
        return None


def _generate_process_map(
    tpl: DocxTemplate,
    steps: list[dict],
    tmp_dir: Path,
) -> Optional[InlineImage]:
    """
    Generate a sequential process map PNG using Pillow and return as InlineImage.
    Draws orange-accented step boxes with downward arrows between them.
    """
    if not steps:
        return None
    try:
        from PIL import Image, ImageDraw, ImageFont

        # Layout
        IMG_W     = 1400
        PADDING   = 50
        BOX_H     = 82
        BOX_R     = 12       # corner radius
        ARROW_H   = 34
        CIRCLE_R  = 24
        HEADER_H  = 72

        # Palette
        ORANGE      = (232, 92, 26)
        LIGHT_GREY  = (245, 245, 245)
        BORDER      = (204, 204, 204)
        TEXT_DARK   = (26, 26, 26)
        WHITE       = (255, 255, 255)

        n = len(steps)
        total_h = HEADER_H + PADDING + n * BOX_H + (n - 1) * ARROW_H + PADDING

        img  = Image.new("RGB", (IMG_W, total_h), WHITE)
        draw = ImageDraw.Draw(img)

        # Fonts (fall back to default pixel font if DejaVu not present)
        try:
            fnt_head = ImageFont.truetype(
                "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 24
            )
            fnt_body = ImageFont.truetype(
                "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 19
            )
            fnt_num  = ImageFont.truetype(
                "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 21
            )
        except Exception:
            fnt_head = fnt_body = fnt_num = ImageFont.load_default()

        # Header bar
        draw.rectangle([(0, 0), (IMG_W, HEADER_H)], fill=ORANGE)
        draw.text((PADDING, HEADER_H // 2 - 14), "Process Flow", font=fnt_head, fill=WHITE)

        y      = HEADER_H + PADDING
        box_x1 = PADDING
        box_x2 = IMG_W - PADDING
        mid_x  = IMG_W // 2

        for i, step in enumerate(steps):
            is_last = i == n - 1

            # Step box
            draw.rounded_rectangle(
                [(box_x1, y), (box_x2, y + BOX_H)],
                radius=BOX_R,
                fill=LIGHT_GREY,
                outline=ORANGE if is_last else BORDER,
                width=2,
            )

            # Numbered circle on the left
            cx = box_x1 + PADDING // 2 + CIRCLE_R
            cy = y + BOX_H // 2
            draw.ellipse(
                [(cx - CIRCLE_R, cy - CIRCLE_R), (cx + CIRCLE_R, cy + CIRCLE_R)],
                fill=ORANGE,
            )
            num_str = str(step.get("sequence", i + 1))
            draw.text((cx - CIRCLE_R // 2 - 1, cy - CIRCLE_R // 2 + 1), num_str, font=fnt_num, fill=WHITE)

            # Step title
            title    = step.get("title", "")
            title    = title[:72] + ("…" if len(title) > 72 else "")
            text_x   = cx + CIRCLE_R + 16
            text_y   = y + BOX_H // 2 - 12
            draw.text((text_x, text_y), title, font=fnt_body, fill=TEXT_DARK)

            y += BOX_H

            # Arrow between steps
            if not is_last:
                arrow_tip = y + ARROW_H
                draw.line([(mid_x, y), (mid_x, arrow_tip - 10)], fill=BORDER, width=2)
                draw.polygon(
                    [(mid_x - 9, arrow_tip - 10), (mid_x + 9, arrow_tip - 10), (mid_x, arrow_tip)],
                    fill=BORDER,
                )
                y += ARROW_H

        map_path = tmp_dir / "process_map.jpg"
        img.save(str(map_path), "JPEG", quality=92, optimize=True)
        return InlineImage(tpl, str(map_path), width=Inches(5.5))

    except Exception as exc:
        logger.warning("Could not generate process map: %s", exc)
        return None


def _hex_to_rgb(hex_color: str) -> tuple[int, int, int]:
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))  # type: ignore[return-value]


def _generate_swimlane_map(
    tpl: DocxTemplate,
    config: dict,
    steps: list[dict],
    tmp_dir: Path,
) -> Optional[InlineImage]:
    """
    Generate a swim-lane process map PNG from process_map_config.
    Lanes are vertical columns; steps flow top-to-bottom with cross-lane arrows.
    """
    if not steps or not config:
        return None
    try:
        from PIL import Image, ImageDraw, ImageFont

        lanes = config.get("lanes", [])
        assignments = config.get("assignments", [])
        if not lanes or not assignments:
            return None

        step_by_id = {s.get("id"): s for s in steps}
        lane_idx = {l["id"]: i for i, l in enumerate(lanes)}

        LANE_W   = 300
        ROW_H    = 110
        BOX_W    = 260
        BOX_H    = 64
        HEADER_H = 58
        MARGIN   = 20
        CIRCLE_R = 18

        n_lanes = len(lanes)
        n_rows  = len(assignments)

        IMG_W = MARGIN + n_lanes * LANE_W + MARGIN
        IMG_H = MARGIN + HEADER_H + n_rows * ROW_H + MARGIN

        WHITE  = (255, 255, 255)
        LIGHT  = (248, 250, 252)
        ALT    = (241, 245, 249)
        BORDER = (203, 213, 225)
        DARK   = (15, 23, 42)
        ARROW  = (148, 163, 184)
        AMBER  = (217, 119, 6)
        CREAM  = (255, 251, 235)

        img  = Image.new("RGB", (IMG_W, IMG_H), WHITE)
        draw = ImageDraw.Draw(img)

        try:
            fnt_head = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 18)
            fnt_body = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 15)
            fnt_num  = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 14)
            fnt_small = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 13)
        except Exception:
            fnt_head = fnt_body = fnt_num = fnt_small = ImageFont.load_default()

        def text_size(text: str, font) -> tuple[int, int]:
            """Return (width, height) of rendered text using textbbox."""
            try:
                bb = draw.textbbox((0, 0), text, font=font)
                return bb[2] - bb[0], bb[3] - bb[1]
            except Exception:
                return len(text) * 8, 16

        def draw_centered(text: str, cx: int, cy: int, font, fill):
            """Draw text centered at (cx, cy)."""
            w, h = text_size(text, font)
            draw.text((cx - w // 2, cy - h // 2), text, font=font, fill=fill)

        def wrap_title(title: str, max_w: int, font) -> list[str]:
            """Wrap title to fit within max_w pixels, max 2 lines."""
            words = title.split()
            lines: list[str] = []
            current = ""
            for word in words:
                test = (current + " " + word).strip()
                w, _ = text_size(test, font)
                if w > max_w and current:
                    lines.append(current)
                    current = word
                    if len(lines) >= 2:
                        break
                else:
                    current = test
            if current and len(lines) < 2:
                lines.append(current)
            # Truncate last line if still too wide
            if lines and text_size(lines[-1], font)[0] > max_w:
                while lines[-1] and text_size(lines[-1] + "…", font)[0] > max_w:
                    lines[-1] = lines[-1][:-1]
                lines[-1] += "…"
            return lines or [""]

        # ── Lane backgrounds + headers ────────────────────────────────────────
        for i, lane in enumerate(lanes):
            lx = MARGIN + i * LANE_W
            ly = MARGIN
            bg = LIGHT if i % 2 == 0 else ALT
            draw.rectangle([(lx, ly), (lx + LANE_W, ly + HEADER_H + n_rows * ROW_H)], fill=bg)
            color_rgb = _hex_to_rgb(lane.get("color", "#3B82F6"))
            draw.rectangle([(lx, ly), (lx + LANE_W, ly + HEADER_H)], fill=color_rgb)
            name = lane.get("name", f"Lane {i + 1}")
            draw_centered(name, lx + LANE_W // 2, ly + HEADER_H // 2, fnt_head, WHITE)
            if i > 0:
                draw.line([(lx, MARGIN), (lx, IMG_H - MARGIN)], fill=BORDER, width=1)

        draw.rectangle([(MARGIN, MARGIN), (IMG_W - MARGIN, IMG_H - MARGIN)], outline=BORDER, width=2)

        def box_center(row: int, lane_id: str) -> tuple[int, int]:
            li = lane_idx.get(lane_id, 0)
            cx = MARGIN + li * LANE_W + LANE_W // 2
            cy = MARGIN + HEADER_H + row * ROW_H + ROW_H // 2
            return cx, cy

        # ── Arrows (drawn behind boxes) ───────────────────────────────────────
        for i, asgn in enumerate(assignments[:-1]):
            next_asgn = assignments[i + 1]
            x1, y1 = box_center(i, asgn["lane_id"])
            x2, y2 = box_center(i + 1, next_asgn["lane_id"])
            is_decision_from = asgn.get("is_decision", False)
            is_decision_to   = next_asgn.get("is_decision", False)
            half_from = (BOX_H // 2 + 8) if is_decision_from else BOX_H // 2
            half_to   = (BOX_H // 2 + 8) if is_decision_to   else BOX_H // 2

            ay_from = y1 + half_from
            ay_to   = y2 - half_to - 4
            mid_y   = y1 + ROW_H // 2

            if x1 == x2:
                draw.line([(x1, ay_from), (x2, ay_to)], fill=ARROW, width=2)
            else:
                draw.line([(x1, ay_from), (x1, mid_y)], fill=ARROW, width=2)
                draw.line([(x1, mid_y), (x2, mid_y)], fill=ARROW, width=2)
                draw.line([(x2, mid_y), (x2, ay_to)], fill=ARROW, width=2)

            draw.polygon([(x2 - 7, ay_to), (x2 + 7, ay_to), (x2, ay_to + 12)], fill=ARROW)

        # ── Step boxes / diamonds ─────────────────────────────────────────────
        for i, asgn in enumerate(assignments):
            step = step_by_id.get(asgn.get("step_id"), {})
            cx, cy = box_center(i, asgn["lane_id"])
            lane   = lanes[lane_idx.get(asgn["lane_id"], 0)]
            color_rgb = _hex_to_rgb(lane.get("color", "#3B82F6"))
            seq_num   = step.get("sequence", i + 1)
            raw_title = step.get("title") or ""

            if asgn.get("is_decision"):
                # Diamond — larger to fit text
                hw, hh = BOX_W // 2, BOX_H // 2 + 10
                draw.polygon([(cx, cy - hh), (cx + hw, cy), (cx, cy + hh), (cx - hw, cy)], fill=CREAM, outline=AMBER, width=2)
                # Sequence label top-left of diamond
                seq_label = f"{seq_num}."
                draw_centered(seq_label, cx - hw // 2, cy - hh // 2 + 4, fnt_num, AMBER)
                # Title wrapped inside diamond (max ~BOX_W - 60px safe inner area)
                title_lines = wrap_title(raw_title, BOX_W - 60, fnt_small)
                line_h = 16
                total_h = len(title_lines) * line_h
                start_y = cy - total_h // 2
                for li, line in enumerate(title_lines):
                    draw_centered(line, cx, start_y + li * line_h, fnt_small, DARK)
            else:
                # Rounded rectangle
                bx = cx - BOX_W // 2
                by = cy - BOX_H // 2
                draw.rounded_rectangle([(bx, by), (bx + BOX_W, by + BOX_H)], radius=8, fill=WHITE, outline=color_rgb, width=2)
                # Number circle
                ccx = bx + 12 + CIRCLE_R
                draw.ellipse([(ccx - CIRCLE_R, cy - CIRCLE_R), (ccx + CIRCLE_R, cy + CIRCLE_R)], fill=color_rgb)
                draw_centered(str(seq_num), ccx, cy, fnt_num, WHITE)
                # Title — wrap to fit in box width minus circle area
                title_area_w = BOX_W - (12 + CIRCLE_R * 2 + 12)
                title_lines = wrap_title(raw_title, title_area_w, fnt_body)
                line_h = 18
                total_h = len(title_lines) * line_h
                tx = ccx + CIRCLE_R + 10
                start_y = cy - total_h // 2
                for li, line in enumerate(title_lines):
                    draw.text((tx, start_y + li * line_h), line, font=fnt_body, fill=DARK)

        map_path = tmp_dir / "process_map_swimlane.jpg"
        img.save(str(map_path), "JPEG", quality=92, optimize=True)
        return InlineImage(tpl, str(map_path), width=Inches(6.5))

    except Exception as exc:
        logger.warning("Could not generate swimlane map: %s", exc)
        return _generate_process_map(tpl, steps, tmp_dir)


def _download_confirmed_map(
    tpl: DocxTemplate,
    url: str,
    tmp_dir: Path,
    sas_token: str = "",
) -> Optional[InlineImage]:
    """Download the user-uploaded confirmed process map PNG and embed it in the document."""
    try:
        full_url = f"{url}?{sas_token}" if sas_token and "?" not in url else url
        resp = requests.get(full_url, timeout=30)
        resp.raise_for_status()
        from PIL import Image as PILImage
        import io as _io
        pmap = PILImage.open(_io.BytesIO(resp.content)).convert("RGB")
        map_path = tmp_dir / "process_map_confirmed.jpg"
        pmap.save(str(map_path), "JPEG", quality=92, optimize=True)
        return InlineImage(tpl, str(map_path), width=Inches(6.5))
    except Exception as exc:
        logger.warning("Could not download confirmed process map from %s: %s", url, exc)
        return None


def _convert_to_pdf(docx_path: Path, output_dir: Path) -> Path:
    """Convert a .docx to .pdf using LibreOffice headless."""
    cmd = [
        "libreoffice",
        "--headless",
        "--convert-to", "pdf",
        "--outdir", str(output_dir),
        str(docx_path),
    ]
    result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice conversion failed: {result.stderr[-500:]}")

    pdf_path = output_dir / docx_path.with_suffix(".pdf").name
    if not pdf_path.exists():
        raise RuntimeError(f"LibreOffice ran but PDF not found at {pdf_path}")
    return pdf_path


def _upload_blob(local_path: Path, sas_url: str, content_type: str, max_retries: int = 3) -> None:
    """PUT a file to Azure Blob Storage using a SAS URL with retry on connection errors."""
    file_size = local_path.stat().st_size
    for attempt in range(max_retries):
        try:
            with open(local_path, "rb") as f:
                resp = requests.put(
                    sas_url,
                    data=f,
                    headers={
                        "x-ms-blob-type": "BlockBlob",
                        "Content-Type": content_type,
                        "Content-Length": str(file_size),
                    },
                    timeout=300,
                )
            resp.raise_for_status()
            return
        except (requests.exceptions.ConnectionError, requests.exceptions.Timeout) as e:
            if attempt < max_retries - 1:
                logger.warning("Blob upload failed (attempt %d/%d): %s — retrying", attempt + 1, max_retries, e)
                time.sleep(5 * (attempt + 1))
            else:
                raise
