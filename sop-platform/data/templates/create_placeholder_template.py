"""
Generates sop_template.docx — Infomate-branded SOP template.
Run once to regenerate the template file:
    pip install python-docx
    python data/templates/create_placeholder_template.py

Context variables consumed by the template:
    sop_title, client_name, process_name, meeting_date,
    generated_date, step_count,
    sections_pre  (list of {section_title, content_text})  -- display_order < 50
    sections_post (list of {section_title, content_text})  -- display_order >= 50
    process_map   (InlineImage | None)
    steps         (list of {sequence, title, description,
                             sub_steps, screenshot, callouts})
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


ORANGE = RGBColor(0xE8, 0x5C, 0x1A)
DARK   = RGBColor(0x1A, 0x1A, 0x1A)
GREY   = RGBColor(0x88, 0x88, 0x88)


def _orange_heading(doc, text: str, level: int = 2):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = ORANGE
    return p


def _meta_table(doc, rows: list[tuple[str, str]]):
    """2-column bold-label / value table for cover metadata."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (label, val) in enumerate(rows):
        cells = table.rows[i].cells
        cells[0].text = label
        if cells[0].paragraphs[0].runs:
            cells[0].paragraphs[0].runs[0].bold = True
        cells[1].text = val
    doc.add_paragraph()  # spacer after table


def create():
    doc = Document()

    # ── COVER PAGE ───────────────────────────────────────────────────────────────
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover.add_run("{{ sop_title }}")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = ORANGE

    doc.add_paragraph()  # spacer

    _meta_table(doc, [
        ("Client",       "{{ client_name }}"),
        ("Process",      "{{ process_name }}"),
        ("Meeting Date", "{{ meeting_date }}"),
        ("Generated",    "{{ generated_date }}"),
        ("Total Steps",  "{{ step_count }}"),
    ])

    doc.add_page_break()

    # ── TABLE OF CONTENTS ────────────────────────────────────────────────────────
    _orange_heading(doc, "Table of Contents", level=1)

    # Each TOC entry on its own line — use {% %} without dashes to preserve paragraph breaks
    doc.add_paragraph("{% for entry in toc_entries %}")

    toc_p = doc.add_paragraph()
    toc_p.paragraph_format.space_after = Pt(2)
    toc_p.paragraph_format.space_before = Pt(2)
    toc_run = toc_p.add_run("{{ '    ' if entry.indent else '' }}{{ entry.title }}")
    toc_run.font.size = Pt(11)

    doc.add_paragraph("{% endfor %}")
    doc.add_page_break()

    # ── PRE-PROCEDURE SECTIONS (display_order < 50) ──────────────────────────────
    doc.add_paragraph("{%- for section in sections_pre %}")
    _orange_heading(doc, "{{ section.section_title }}", level=2)
    doc.add_paragraph("{{ section.content_text | default('') }}")
    doc.add_paragraph("{%- endfor %}")

    doc.add_page_break()

    # ── PROCESS MAP ──────────────────────────────────────────────────────────────
    _orange_heading(doc, "Process Map", level=1)
    doc.add_paragraph(
        "{%- if process_map %}{{ process_map }}{%- else %}(Process map unavailable){%- endif %}"
    )

    doc.add_page_break()

    # ── DETAILED PROCEDURE ───────────────────────────────────────────────────────
    _orange_heading(doc, "Detailed Procedure", level=1)

    doc.add_paragraph("{%- for step in steps %}")

    step_h = doc.add_heading("Step {{ step.sequence }}: {{ step.title }}", level=3)
    for run in step_h.runs:
        run.font.color.rgb = DARK

    doc.add_paragraph("{{ step.description | default('') }}")

    # Sub-steps as bullet list
    doc.add_paragraph("{%- for sub in step.sub_steps %}")
    sub_p = doc.add_paragraph(style="List Bullet")
    sub_p.add_run("{{ sub }}")
    doc.add_paragraph("{%- endfor %}")

    # Annotated screenshot (InlineImage)
    doc.add_paragraph(
        "{%- if step.screenshot %}{{ step.screenshot }}{%- endif %}"
    )

    # Callout legend
    doc.add_paragraph("{%- if step.callouts %}")
    legend_label = doc.add_paragraph()
    r = legend_label.add_run("Callout References")
    r.bold = True
    r.font.size = Pt(10)
    doc.add_paragraph("{%- for callout in step.callouts %}")
    doc.add_paragraph("{{ callout.callout_number }}. {{ callout.label }}")
    doc.add_paragraph("{%- endfor %}")
    doc.add_paragraph("{%- endif %}")

    doc.add_paragraph("{%- endfor %}")  # end steps loop

    doc.add_page_break()

    # ── POST-PROCEDURE SECTIONS (display_order >= 50) ────────────────────────────
    doc.add_paragraph("{%- for section in sections_post %}")
    _orange_heading(doc, "{{ section.section_title }}", level=2)
    doc.add_paragraph("{{ section.content_text | default('') }}")
    doc.add_paragraph("{%- endfor %}")

    out = Path(__file__).parent / "sop_template.docx"
    doc.save(out)
    print(f"Template created: {out}")


if __name__ == "__main__":
    create()
